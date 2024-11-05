import json
from docx import Document


class DocxPlaceholderReplacer:
    def __init__(self, demo_file_path, user_input_path, output_path):
        self.demo_file = demo_file_path
        self.user_input_path = user_input_path
        self.output_path = output_path
        self.user_input = self.load_user_input()
        self.document = self.load_demo_file_format()

    def load_user_input(self):
        try:
            with open(self.user_input_path, 'r', encoding='utf-8') as file:
                return json.load(file)
        except Exception as e:
            print(f"Error loading user input: {e}")
            return {}

    def load_demo_file_format(self):
        try:
            return Document(self.demo_file)
        except Exception as e:
            print(f"Error loading template: {e}")
            return None

    def replace_placeholders(self):
        if self.document is None:
            return

        for para in self.document.paragraphs:
            full_text = ''.join(run.text for run in para.runs)

            updated_text = full_text
            for key, data in self.user_input.items():
                placeholder = f"[[{key}]]"
                if placeholder in updated_text:
                    updated_text = updated_text.replace(placeholder, data['value'])

            if updated_text == full_text:
                continue

            for run in para.runs:
                run.clear()

            run = para.add_run(updated_text)

            for key, data in self.user_input.items():
                if data.get("label", "").lower() == "bold":
                    bold_placeholder = f"[[{key}]]"
                    if bold_placeholder in full_text:
                        run.bold = True

    def save_updated_document(self):
        try:
            self.document.save(self.output_path)
            print(f"Document updated and saved to {self.output_path}")
        except Exception as e:
            print(f"Error saving document: {e}")

    def process(self):
        self.replace_placeholders()
        self.save_updated_document()


demo_file_path = 'Test document.docx'
# user_input_path = 'user_input_for_test_document.txt'
user_input_path = input("Enter user input file path: ")
output_path = 'updated_document.docx'

replacer = DocxPlaceholderReplacer(demo_file_path, user_input_path, output_path)
replacer.process()
